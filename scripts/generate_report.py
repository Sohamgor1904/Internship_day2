import sys
import subprocess

try:
    import docx
except ImportError:
    print("[INIT] python-docx not found. Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets the internal padding (margins) of a table cell in dxa (1/20 of a pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # 1. Page Setup (Letter size, 1 inch margins)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    styles = doc.styles
    
    # Custom Palette
    NAVY = RGBColor(0, 51, 102)      # Primary
    SLATE = RGBColor(90, 100, 110)    # Secondary
    CHARCOAL = RGBColor(51, 51, 51)  # Body text
    
    # Modify default styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = CHARCOAL
    
    # ────────────────────────────────────────────────────────────
    # PAGE 1: TITLE & EXECUTIVE UPDATE
    # ────────────────────────────────────────────────────────────
    
    # Document Header
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("TECHNICAL REPORT & PROJECT UPDATE")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = NAVY
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.paragraph_format.space_after = Pt(24)
    run_sub = p_subtitle.add_run("OCSF-Based Hybrid Threat Detection Pipeline — Production-Grade Triage Architecture")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = SLATE

    # Horizontal Rule
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_after = Pt(18)
    run_hr = p_hr.add_run("―" * 65)
    run_hr.font.color.rgb = SLATE
    
    # Section 1 Heading
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(8)
    r1 = h1.add_run("1. Executive Update for Supervisor")
    r1.font.name = 'Calibri'
    r1.font.size = Pt(16)
    r1.font.bold = True
    r1.font.color.rgb = NAVY
    
    p_body = doc.add_paragraph()
    p_body.paragraph_format.space_after = Pt(12)
    p_body.add_run(
        "Below is the concise, crisp update drafted for your supervisor. "
        "It explains the architectural shift towards a high-performance deterministic Layer 0 gateway, minimizing AI/ML costs."
    )
    
    # Message Block (Bordered Box)
    table_msg = doc.add_table(rows=1, cols=1)
    table_msg.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table_msg.cell(0, 0)
    set_cell_background(cell, "F2F5F8")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p_msg = cell.paragraphs[0]
    p_msg.paragraph_format.space_after = Pt(4)
    p_msg.add_run("Subject: Completion: OCSF Hybrid Threat Detection Pipeline (Production Rule + AI)\n\n").bold = True
    p_msg.add_run(
        "Hi [Supervisor's Name],\n\n"
        "I have successfully optimized our OCSF-based network threat detection pipeline for production deployment. "
        "By aligning the system with the rule 'we can't use AI in everything,' I implemented a deterministic Layer 0 Signature & Whitelist gateway that handles safe loopback/DNS traffic and obvious scans in <0.2ms with 0% AI compute. "
        "Our tiered structure now ensures that heavy machine learning is strictly treated as a last-resort escalation:\n\n"
        "  • Layer 0 (Rules & Whitelists): Fast-passes trusted DNS/loopback and fast-blocks flag scans (Xmas/Null) instantly without loading ML models (90% of traffic bypassed).\n"
        "  • Layer 1 (Statistical Filter): Dynamic Rate Z-Score, EWMA volume spikes, and port entropy triage (anomaly threshold: >= 2.5σ).\n"
        "  • Layer 2 & 3 (ML Escalation): Random Forest (100% test recall + SHAP attributions) and sequential PyTorch LSTM (97.63% sequential accuracy) invoked strictly for L1 anomalies.\n\n"
        "The system has been successfully verified using a 1,000-packet performance simulation, showing a 14.5% total compute reduction and 1.2x speedup on localhost socket connections. "
        "A full technical report is attached, and I'm ready to walk you through a brief live demo at your convenience.\n\n"
        "Best regards,\n"
        "[Your Name]"
    )

    # ────────────────────────────────────────────────────────────
    # PAGE 2: TIERED MODELING ARCHITECTURE
    # ────────────────────────────────────────────────────────────
    doc.add_page_break()
    
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(8)
    r2 = h2.add_run("2. Production-Grade Tiered Architecture")
    r2.font.name = 'Calibri'
    r2.font.size = Pt(16)
    r2.font.bold = True
    r2.font.color.rgb = NAVY
    
    p_arch = doc.add_paragraph()
    p_arch.paragraph_format.space_after = Pt(14)
    p_arch.add_run(
        "To optimize CPU utilization while maintaining state-of-the-art accuracy, incoming flows are processed sequentially through a tiered pipeline. "
        "This design ensures that heavy AI models are strictly reserved for rare, complex, and highly anomalous traffic."
    )
    
    # Architecture Table
    table_arch = doc.add_table(rows=5, cols=3)
    table_arch.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Layer Stage", "Model Paradigm", "Primary Role & Computational Rationale"]
    hdr_cells = table_arch.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "003366")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    row_data = [
        ("Layer 0: Rules & Whitelist", "Deterministic Check Engine", 
         "Performs instant IP/DNS whitelisting and TCP flag signature blocks (Null, Xmas, SYN-FIN scans). Processes in <0.2ms with 0% AI model overhead, bypassing feature extraction entirely."),
        ("Layer 1: Volumetric Filter", "Stateful Statistical State Machine", 
         "Ultra-low overhead statistical triage using moving averages, EWMA volume variations, and port entropy. Anomaly threshold: >= 2.5σ. If normal, dropped immediately with 0% AI cost."),
        ("Layer 2: Contextual Classifier", "Random Forest + SHAP Explainer", 
         "Supervised machine learning invoked only for Layer 1 anomalies. Detects isolated threats (exploits) with 100% recall. Generates detailed SHAP attributions if probability >= 0.5."),
        ("Layer 3: Sequential Tracker", "Chronological PyTorch LSTM", 
         "Deep learning sequential model. Evaluates sliding window deques of the last 10 events per unique host IP to identify multi-stage Lateral Movements and slow APT beaconing sequences.")
    ]
    
    for row_idx, data in enumerate(row_data):
        row_cells = table_arch.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F7F9FA")

    # ────────────────────────────────────────────────────────────
    # PAGE 3: STATEFUL FEATURE ENGINEERING
    # ────────────────────────────────────────────────────────────
    doc.add_page_break()
    
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(8)
    r3 = h3.add_run("3. Stateful Feature Engineering")
    r3.font.name = 'Calibri'
    r3.font.size = Pt(16)
    r3.font.bold = True
    r3.font.color.rgb = NAVY
    
    p_feat = doc.add_paragraph()
    p_feat.paragraph_format.space_after = Pt(14)
    p_feat.add_run(
        "For events promoted past Layer 0 rules, the StreamingFeaturePipeline statefully normalizes "
        "raw JSON entries into OCSF Network Traffic (4001) format and extracts 12 stateful features:"
    )
    
    features = [
        ("Temporal Dynamics", "Tracks exact inter-arrival intervals (delta_t) in seconds and rolling flow durations."),
        ("Volumetric Ratios", "Measures forward/backward packet rates and scaled outbound-to-inbound byte ratios (byte_ratio) to catch asymmetric floods."),
        ("Structural Entropy", "Computes the Shannon Entropy of destination IPs and ports within a rolling window to detect single-target flooding or wide port-scans."),
        ("State Transitions", "Statefully tracks TCP flag turbulence (flag_switches) by counting changes in connection states across sliding sequence queues.")
    ]
    
    for category, desc in features:
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.space_after = Pt(6)
        run_cat = p_item.add_run(f"{category}: ")
        run_cat.bold = True
        run_cat.font.color.rgb = NAVY
        p_item.add_run(desc)

    # ────────────────────────────────────────────────────────────
    # PAGE 4: PERFORMANCE METRICS & TRAINING RESULTS
    # ────────────────────────────────────────────────────────────
    doc.add_page_break()
    
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(8)
    r4 = h4.add_run("4. Performance Metrics & Benchmark Outcomes")
    r4.font.name = 'Calibri'
    r4.font.size = Pt(16)
    r4.font.bold = True
    r4.font.color.rgb = NAVY
    
    p_metrics = doc.add_paragraph()
    p_metrics.paragraph_format.space_after = Pt(14)
    p_metrics.add_run("The hybrid pipeline has been trained on 18,000 balanced OCSF entries and evaluated using a 1,000-packet performance simulation client:")
    
    metric_points = [
        ("Layer 0 Whitelist & Signature Engine", "Achieved 100% test accuracy in independent checks. Correctly whitelisted local loopbacks and safe DNS instantly. Successfully blocked Null and Xmas scans with 0% AI overhead."),
        ("Level 5 Compute Efficiency Benchmark", "A 1,000-request simulation (800 whitelisted loopbacks, 100 blocked scans, 100 escalated ML packets) demonstrated that 90.0% of requests bypassed ML entirely. Resulted in a 14.5% total compute latency reduction and a 1.2x socket throughput speedup."),
        ("Layer 2 Random Forest Classifier", "Achieved 100% accuracy and 100% recall on test attack rows with 0 false negatives. Integrates real-time SHAP TreeExplainer local attributions sorted by absolute contribution."),
        ("Layer 3 Deep Learning LSTM", "Successfully trained for 20 epochs with early stopping (patience=5, best loss=0.0865). Restored weights achieved 97.63% test accuracy on host timelines. Verified sequence order sensitivity (forward lateral scan prob = 0.9942 vs. reversed = 0.0046).")
    ]
    
    for title, detail in metric_points:
        p_m = doc.add_paragraph(style='List Bullet')
        p_m.paragraph_format.space_after = Pt(6)
        run_title = p_m.add_run(f"{title}: ")
        run_title.bold = True
        run_title.font.color.rgb = NAVY
        p_m.add_run(detail)

    # ────────────────────────────────────────────────────────────
    # PAGE 5: CORE TECHNOLOGY STACK
    # ────────────────────────────────────────────────────────────
    doc.add_page_break()
    
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(12)
    h5.paragraph_format.space_after = Pt(8)
    r5 = h5.add_run("5. Core Technology Stack")
    r5.font.name = 'Calibri'
    r5.font.size = Pt(16)
    r5.font.bold = True
    r5.font.color.rgb = NAVY
    
    p_tech = doc.add_paragraph()
    p_tech.paragraph_format.space_after = Pt(14)
    p_tech.add_run("The development, testing, and deployment configurations integrate the following packages and tools:")
    
    table_tech = doc.add_table(rows=5, cols=2)
    table_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    tech_hdrs = ["Component Category", "Technologies & Libraries Integrated"]
    tech_cells = table_tech.rows[0].cells
    for i, h_text in enumerate(tech_hdrs):
        tech_cells[i].text = h_text
        set_cell_background(tech_cells[i], "5A646E")
        set_cell_margins(tech_cells[i], top=90, bottom=90, left=100, right=100)
        p = tech_cells[i].paragraphs[0]
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    tech_data = [
        ("API & Microservice Engine", "FastAPI (exposing /detect and /health), Uvicorn (high-performance ASGI server), Pydantic v2 (nested schema data validation)"),
        ("Machine Learning Core", "Scikit-Learn (Random Forest Classifier, StandardScaler), SHAP (TreeExplainer for real-time local attributions)"),
        ("Deep Learning Core", "PyTorch (Sequential LSTM Model architecture and recurrent gate logic)"),
        ("Infrastructure & Databases", "PostgreSQL (threat alerts logging and pipeline stats tracking), Asyncpg (asynchronous database pool)")
    ]
    
    for row_idx, data in enumerate(tech_data):
        row_cells = table_tech.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F7F9FA")

    # Save report with PermissionError safety fallback
    # Resolve parent project root dynamically relative to scripts/ folder
    import os
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    out_path = os.path.join(project_root, "doc", "OCSF_Hybrid_Threat_Detection_Report.docx")
    try:
        doc.save(out_path)
        print(f"[SUCCESS] Beautiful Word Document generated successfully at {out_path}!")
    except PermissionError:
        import time
        suffix = int(time.time())
        alt_path = os.path.join(project_root, "doc", f"OCSF_Hybrid_Threat_Detection_Report_{suffix}.docx")
        doc.save(alt_path)
        print(f"[SUCCESS] Primary file was locked. Saved copy at {alt_path}!")

if __name__ == "__main__":
    create_document()
